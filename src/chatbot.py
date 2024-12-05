import streamlit as st
from langchain_openai import ChatOpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from dotenv import load_dotenv
from tempfile import NamedTemporaryFile
from langchain.schema import AIMessage, HumanMessage, SystemMessage
import re
from langchain.document_loaders import PyPDFium2Loader as PyPDFLoader


# Cargar variables desde el archivo .env
load_dotenv()

# Función para limpiar texto de caracteres de Markdown y HTML
def clean_text(text):
    """Elimina caracteres de Markdown y HTML del texto para hacerlo más legible."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # Eliminar negritas de Markdown
    text = re.sub(r"\*(.*?)\*", r"\1", text)      # Eliminar itálicas de Markdown
    text = re.sub(r"<br>", "\n", text)           # Reemplazar <br> por saltos de línea
    text = re.sub(r"`(.*?)`", r"\1", text)       # Eliminar backticks
    return text.strip()

# Crear encabezado para secciones
def add_section_header(doc, title):
    """Agrega un encabezado de sección al documento."""
    header = doc.add_paragraph()
    run = header.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Crear documento Word con tablas y texto limpio
def create_word_document_with_clean_formatting(response_text):
    """Crea un archivo Word limpio basado en el texto procesado."""
    doc = Document()
    doc.add_heading("Resumen generado por Solutia", level=1)

    lines = response_text.split("\n")
    table_lines = []
    is_table = False

    for line in lines:
        line = clean_text(line)

        # Identificar inicio y fin de tablas
        if line.startswith("|") and line.endswith("|"):
            is_table = True
            table_lines.append(line)
        elif is_table and line.strip() == "":
            is_table = False
            if table_lines:
                add_table_to_document(doc, table_lines)
                table_lines = []
        elif is_table:
            table_lines.append(line)
        else:
            # Agregar texto normal o encabezados
            if line.startswith("##") or line.startswith("**"):
                add_section_header(doc, line.replace("##", "").strip())
            else:
                para = doc.add_paragraph()
                para.add_run(line).font.size = Pt(11)

    # Procesar cualquier tabla restante
    if table_lines:
        add_table_to_document(doc, table_lines)

    # Guardar el documento
    doc_path = "resumen_generado.docx"
    doc.save(doc_path)
    return doc_path

def add_table_to_document(doc, table_lines):
    """Convierte líneas con formato de tabla en una tabla legible dentro del documento Word."""
    rows = [line.strip("|").split("|") for line in table_lines]
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"

    # Encabezados de la tabla
    header_cells = table.rows[0].cells
    for i, header in enumerate(rows[0]):
        header_cells[i].text = clean_text(header.strip())
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas de datos
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            cells[i].text = clean_text(cell.strip())

# Configuración del modelo
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0,
    api_key=os.getenv("OPENAI_API_KEY")
)

if "messages" not in st.session_state:
    st.session_state.messages = []

# Función para extraer texto del PDF
def extract_text_with_langchain(uploaded_file):
    """Extrae texto de un archivo PDF utilizando LangChain PyPDFLoader."""
    with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_file_path = temp_file.name

    try:
        loader = PyPDFLoader(temp_file_path)
        documents = loader.load()
        text = "\n".join([doc.page_content for doc in documents])
        if not text.strip():
            st.error(f"El archivo {uploaded_file.name} no contiene texto o no pudo ser leído.")
        return text
    finally:
        os.remove(temp_file_path)

# Procesar texto del documento completo
def process_full_document(text, model, task="Resumen del documento"):
    """Procesa todo el texto del documento y genera un resumen profesional y limpio."""
    system_message = SystemMessage(
        content="""Eres un chatbot AI útil, te llamas Solutia, tienes que hablar como si fueras un asistente, nunca reveles que eres un chatbot.
Rol
Eres un experto en análisis de documentos técnicos y generación de resúmenes estructurados para licitaciones públicas. Tienes años de experiencia analizando pliegos de prescripciones técnicas (PPT) y pliegos de cláusulas administrativas particulares (PCAP) de concursos públicos. Posees habilidades avanzadas para extraer y sintetizar la información más relevante de estos documentos, asegurando que los responsables de buscar licitaciones puedan tomar decisiones rápidamente.

Tarea
Tu objetivo es analizar los documentos adjuntos (PPT, PCAP o ambos) relacionados con un concurso público y generar los siguientes entregables:

Para el PPT:
Resumen en texto corrido (500 palabras):
Proporciona un análisis detallado de las condiciones técnicas del contrato.
Estructura el contenido por temas relevantes (como niveles de servicio, inventario, mantenimiento, etc.).
Asegúrate de incluir un desglose detallado del Acuerdo a niveles y servicios (ANS), identificando:
Indicadores de desempeño.
Penalizaciones en caso de incumplimiento.
Mecanismos de seguimiento y evaluación técnica​1
.
Para el PCAP:
Resumen detallado en formato tabla (1000 palabras) que incluya:
Cómo acreditar la solvencia técnica y financiera: Explica con claridad los medios de acreditación y los requisitos mínimos exigidos por la normativa o el documento​2
.
Importes máximos de licitación por lote (sin IVA): Si aplicable, desglosados por lotes o en su totalidad​2
.
Criterios de adjudicación: Divide esta sección en dos grandes apartados:
A. Criterios Objetivos (Cuantificables mediante fórmulas matemáticas):
Explica cómo se puntúan y qué peso específico tienen.
Proporciona ejemplos claros de cómo se aplican las fórmulas y qué información se debe proporcionar para cumplir estos requisitos​2
.
Ejemplo: Puntos otorgados por reducir el coste o por aumentar los tiempos de garantía.
B. Criterios Subjetivos (Juicios de valor):
Desarrolla ampliamente qué aspectos evalúan y cómo se otorgan los puntos.
Describe si son evaluados por un comité técnico o un grupo de expertos.
Ejemplo: Calidad técnica de la propuesta, metodología de trabajo, mejoras ofrecidas, capacidad de innovación, etc.​2
.
Otros datos relevantes:
Presupuesto base de licitación, valor estimado del contrato y plazos​2
.
Garantías: Provisional y definitiva, especificando importes y condiciones​2
.
Lugar y medios de presentación de ofertas​2
.
Detalles específicos:
Criterios Objetivos y Subjetivos:
Desglosa en detalle cómo cada criterio impacta la valoración final y cuál es su peso en el resultado total.
Proporciona ejemplos concretos de cómo se aplican, en base a las fórmulas o los juicios definidos en el PCAP.
Si falta información en los documentos, indícalo y sugiere cómo podrían completarse con base en normativas aplicables (como la LCSP).
Contexto:
Solutia es una empresa especializada en identificar y preparar documentación para concursos públicos mediante el uso de inteligencia artificial. Este bot tiene como objetivo optimizar la carga de trabajo de los responsables de licitaciones, proporcionando resúmenes estructurados y fáciles de interpretar que les permitan ahorrar tiempo y enfocarse en las oportunidades más relevantes.

Notas Adicionales:
Si algún parámetro clave no aparece en los documentos, debes indicarlo claramente en los resúmenes con una nota específica y referencia a la LCSP o normativas aplicables.
Garantiza que el resumen sea profesional, claro y adecuado para uso interno en la empresa.
Si algún criterio necesita ampliarse, añade ejemplos hipotéticos para hacerlo más claro."""
    )
    user_message = HumanMessage(content=f"Tarea: {task}\n\nTexto del documento:\n{text}")

    try:
        response = model([system_message, user_message])
        return response.content.strip()
    except Exception as e:
        st.error(f"Error al generar el resumen: {str(e)}")
        return "Error al generar el resumen."

# Subir múltiples archivos
uploaded_files = st.file_uploader("Carga uno o más archivos PDF", type=["pdf"], accept_multiple_files=True)

# Inicializar las variables de texto
pcap_text, ppt_text = "", ""

if uploaded_files:
    with st.spinner("Procesando archivos..."):
        for uploaded_file in uploaded_files:
            try:
                file_text = extract_text_with_langchain(uploaded_file)
                if file_text:
                    st.text_area(f"Texto extraído de {uploaded_file.name}", value=file_text, height=200)
                if "PCAP" in uploaded_file.name.upper():
                    pcap_text += file_text
                elif "PPT" in uploaded_file.name.upper():
                    ppt_text += file_text
            except Exception as e:
                st.error(f"Error al procesar el archivo {uploaded_file.name}: {str(e)}")
        st.success("Archivos procesados exitosamente.")

# Chatbot funcionalidad
if prompt := st.chat_input("Escribe tu mensaje..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    if "PPT" in prompt.upper() and ppt_text:
        input_text = ppt_text
        document_type = "PPT"
    elif "PCAP" in prompt.upper() and pcap_text:
        input_text = pcap_text
        document_type = "PCAP"
    else:
        input_text = "No se detectó el tipo de documento solicitado o no se cargó."
        document_type = "Desconocido"

    if input_text != "No se detectó el tipo de documento solicitado o no se cargó.":
        with st.spinner("Generando el resumen..."):
            try:
                processed_text = process_full_document(input_text, llm, task=f"Resumen de {document_type}")
                if processed_text.strip():
                    st.markdown("### Resumen:")
                    st.markdown(processed_text, unsafe_allow_html=True)

                    # Generar y descargar archivo Word
                    doc_path = create_word_document_with_clean_formatting(processed_text)
                    try:
                        with open(doc_path, "rb") as file:
                            st.download_button(
                                label=f"Descargar resumen {document_type} en Word",
                                data=file,
                                file_name=f"resumen_{document_type.lower()}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    finally:
                        if os.path.exists(doc_path):
                            os.remove(doc_path)
                else:
                    st.error("El resumen generado está vacío. Verifica el archivo original.")
            except Exception as e:
                st.error(f"Error al procesar el texto: {str(e)}")
