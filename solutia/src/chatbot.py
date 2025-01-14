import streamlit as st
from langchain_openai import ChatOpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from dotenv import load_dotenv
from tempfile import NamedTemporaryFile
from langchain.schema import AIMessage, HumanMessage, SystemMessage
import re
from langchain.document_loaders import PyPDFium2Loader as PyPDFLoader

# Cargar variables desde el archivo .env
load_dotenv()

# Función para limpiar texto de caracteres de Markdown y HTML
def clean_text(text):
    """Elimina caracteres de Markdown y HTML del texto para hacerlo más legible."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # Eliminar negritas de Markdown
    text = re.sub(r"\*(.*?)\*", r"\1", text)      # Eliminar itálicas de Markdown
    text = re.sub(r"<br>", "\n", text)           # Reemplazar <br> por saltos de línea
    text = re.sub(r"`(.*?)`", r"\1", text)       # Eliminar backticks
    return text.strip()

# Crear encabezado para secciones
def add_section_header(doc, title):
    """Agrega un encabezado de sección al documento."""
    header = doc.add_paragraph()
    run = header.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Crear documento Word con tablas y texto limpio
def create_word_document_with_clean_formatting(response_text):
    """Crea un archivo Word limpio basado en el texto procesado."""
    doc = Document()
    doc.add_heading("Resumen generado por Solutia", level=1)

    lines = response_text.split("\n")
    table_lines = []
    is_table = False

    for line in lines:
        line = clean_text(line)

        # Identificar inicio y fin de tablas
        if line.startswith("|") and line.endswith("|"):
            is_table = True
            table_lines.append(line)
        elif is_table and line.strip() == "":
            is_table = False
            if table_lines:
                add_table_to_document(doc, table_lines)
                table_lines = []
        elif is_table:
            table_lines.append(line)
        else:
            # Agregar texto normal o encabezados
            if line.startswith("##") or line.startswith("**"):
                add_section_header(doc, line.replace("##", "").strip())
            else:
                para = doc.add_paragraph()
                para.add_run(line).font.size = Pt(11)

    # Procesar cualquier tabla restante
    if table_lines:
        add_table_to_document(doc, table_lines)

    # Guardar el documento
    doc_path = "resumen_generado.docx"
    doc.save(doc_path)
    return doc_path

def add_table_to_document(doc, table_lines):
    """Convierte líneas con formato de tabla en una tabla legible dentro del documento Word."""
    rows = [line.strip("|").split("|") for line in table_lines]
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"

    # Encabezados de la tabla
    header_cells = table.rows[0].cells
    for i, header in enumerate(rows[0]):
        header_cells[i].text = clean_text(header.strip())
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas de datos
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            cells[i].text = clean_text(cell.strip())


# Configuración del modelo
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0,
    api_key=os.getenv("OPENAI_API_KEY")
)

if "processed_files" not in st.session_state:
    st.session_state.processed_files = {"ppt": None, "pcap": None}

if "processed_summaries" not in st.session_state:
    st.session_state.processed_summaries = {"ppt": None, "pcap": None}

# Agregar nuevo estado para el último tipo de documento procesado
if "last_processed" not in st.session_state:
    st.session_state.last_processed = None

# Agregar nuevo estado para controlar el orden de visualización
if "display_order" not in st.session_state:
    st.session_state.display_order = []

# Función para extraer texto del PDF
def extract_text_with_langchain(uploaded_file):
    """Extrae texto de un archivo PDF utilizando LangChain PyPDFLoader."""
    with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_file_path = temp_file.name

    try:
        loader = PyPDFLoader(temp_file_path)
        documents = loader.load()
        text = "\n".join([doc.page_content for doc in documents])
        if not text.strip():
            st.error(f"El archivo {uploaded_file.name} no contiene texto o no pudo ser leído.")
        return text
    finally:
        os.remove(temp_file_path)

# Procesar texto del documento completo
def split_text(text, max_chunk_size=50000):
    """Divide el texto en chunks más pequeños."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_size = 0
    
    for word in words:
        # Estimación aproximada: 4 caracteres = 1 token
        word_size = len(word) // 4
        if current_size + word_size > max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_size = word_size
        else:
            current_chunk.append(word)
            current_size += word_size
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def process_full_document(text, model, task="Resumen del documento"):
    """Procesa todo el texto del documento y genera un resumen profesional y limpio."""
    # Dividir el texto en chunks más pequeños
    chunks = split_text(text)
    summaries = []
    
    system_message = SystemMessage(
        content="""Eres un chatbot AI útil, te llamas Solutia, tienes que hablar como si fueras un asistente, nunca reveles que eres un chatbot.
Rol
Eres un experto en análisis de documentos técnicos y generación de resúmenes estructurados para licitaciones públicas. Tienes años de experiencia analizando pliegos de prescripciones técnicas (PPT) y pliegos de cláusulas administrativas particulares (PCAP) de concursos públicos. Posees habilidades avanzadas para extraer y sintetizar la información más relevante de estos documentos, asegurando que los responsables de buscar licitaciones puedan tomar decisiones rápidamente.

Tarea
Tu objetivo es analizar los documentos adjuntos (PPT, PCAP o ambos) relacionados con un concurso público y generar los siguientes entregables:

Para el PPT:
Resumen en texto corrido (500 palabras):
Proporciona un análisis detallado de las condiciones técnicas del contrato.
Estructura el contenido por temas relevantes (como niveles de servicio, inventario, mantenimiento, etc.).
Asegúrate de incluir un desglose detallado del Acuerdo a niveles y servicios (ANS), identificando:
Indicadores de desempeño.
Penalizaciones en caso de incumplimiento.
Mecanismos de seguimiento y evaluación técnica​1
.

Aquí tienes el texto reescrito con la información adicional integrada:

Para el PCAP:

Resumen detallado en formato tabla (1000 palabras) que incluya:

Tipo de contrato (servicio / suministro):

Especificar si el contrato corresponde a un servicio o un suministro, detallando las características particulares de cada opción y su impacto en los requisitos y criterios de valoración.
Tipo de suministro / servicio:

Describir en detalle el tipo específico de suministro o servicio requerido, identificando sus características esenciales y las implicaciones en la ejecución del contrato.
Objeto del contrato:

Definir claramente el objetivo principal del contrato, especificando el alcance, las necesidades que busca cubrir y los resultados esperados.
Cómo acreditar la solvencia técnica y financiera:

Explicar con claridad los medios de acreditación y los requisitos mínimos exigidos por la normativa o el documento.
Incluir detalles sobre las certificaciones necesarias (si procede) que pueden estar incluidas dentro de los criterios de valoración.
Importes máximos de licitación por lote (sin IVA):

Si es aplicable, desglosar los importes por lotes o en su totalidad, especificando los límites económicos establecidos para cada uno de ellos.
Presupuesto del contrato:

Indicar el presupuesto total asignado al contrato, diferenciando entre los importes con y sin IVA si corresponde.
Especificar el valor estimado del contrato teniendo en cuenta su duración y posibles prórrogas.
Criterios de adjudicación:

A. Criterios Objetivos (Cuantificables mediante fórmulas matemáticas):
Explicar cómo se puntúan y qué peso específico tienen cada uno de los criterios objetivos.
Proporcionar ejemplos claros de cómo se aplican las fórmulas y qué información se debe proporcionar para cumplir estos requisitos.
Ejemplo: Puntos otorgados por reducir el coste o por aumentar los tiempos de garantía.
B. Criterios Subjetivos (Juicios de valor):
Desarrollar ampliamente qué aspectos evalúan y cómo se otorgan los puntos.
Describir si son evaluados por un comité técnico o un grupo de expertos.
Ejemplo: Calidad técnica de la propuesta, metodología de trabajo, mejoras ofrecidas, capacidad de innovación, etc.
Criterios de valoración:
Incluir un desglose claro de los criterios específicos aplicables, su ponderación y la relación con el objeto del contrato.
Especificar si las certificaciones influyen en estos criterios y cómo se integran en el proceso de puntuación.
Estudio de la fórmula del precio:

Desarrollar la fórmula matemática utilizada para determinar el precio final de adjudicación.
Explicar con ejemplos su aplicación práctica y los factores que afectan al cálculo final.
Condiciones generales:

Establecer las condiciones básicas que rigen el contrato, incluidas las obligaciones y derechos de las partes involucradas.
Especificar las cláusulas de cumplimiento, penalizaciones por incumplimiento y otros aspectos clave.
Qué margen estimado para cada tipo:

Proporcionar una estimación de los márgenes esperados para cada tipo de suministro o servicio.
Especificar si existen diferencias relevantes entre los distintos lotes o fases del contrato.
Otros datos relevantes:
Presupuesto base de licitación: Indicar el presupuesto inicial propuesto para el contrato.
Valor estimado del contrato: Ofrecer una visión clara del valor total del contrato, teniendo en cuenta su duración y posibles ampliaciones.
Plazos: Detallar los plazos de ejecución, entrega y finalización del contrato, así como las posibles prórrogas.
Garantías:
Garantía Provisional: Especificar los importes y condiciones necesarias para su cumplimiento.
Garantía Definitiva: Desglosar las condiciones, plazos e importes requeridos para esta garantía.
Lugar y medios de presentación de ofertas:
Precisar la dirección física y electrónica para la presentación de ofertas.
Incluir detalles sobre el formato, los plazos y los requisitos adicionales para la presentación válida de las ofertas.
.
Detalles específicos:
Criterios Objetivos y Subjetivos:
Desglosa en detalle cómo cada criterio impacta la valoración final y cuál es su peso en el resultado total.
Proporciona ejemplos concretos de cómo se aplican, en base a las fórmulas o los juicios definidos en el PCAP.
Si falta información en los documentos, indícalo y sugiere cómo podrían completarse con base en normativas aplicables (como la LCSP).
Contexto:
Solutia es una empresa especializada en identificar y preparar documentación para concursos públicos mediante el uso de inteligencia artificial. Este bot tiene como objetivo optimizar la carga de trabajo de los responsables de licitaciones, proporcionando resúmenes estructurados y fáciles de interpretar que les permitan ahorrar tiempo y enfocarse en las oportunidades más relevantes.

Notas Adicionales:
Si algún parámetro clave no aparece en los documentos, debes indicarlo claramente en los resúmenes con una nota específica y referencia a la LCSP o normativas aplicables.
Garantiza que el resumen sea profesional, claro y adecuado para uso interno en la empresa.
Si algún criterio necesita ampliarse, añade ejemplos hipotéticos para hacerlo más claro."""
    )
    try:
        # Procesar cada chunk por separado
        for i, chunk in enumerate(chunks):
            chunk_task = f"{task} (Parte {i+1}/{len(chunks)})"
            user_message = HumanMessage(content=f"Tarea: {chunk_task}\n\nTexto del documento:\n{chunk}")
            
            response = model([system_message, user_message])
            summaries.append(response.content.strip())
        
        # Si hay múltiples chunks, hacer un resumen final
        if len(summaries) > 1:
            final_summary_text = "\n\n".join(summaries)
            user_message = HumanMessage(
                content=f"Tarea: Generar resumen final combinando los siguientes resúmenes parciales:\n\n{final_summary_text}"
            )
            final_response = model([system_message, user_message])
            return final_response.content.strip()
        
        return summaries[0]
    
    except Exception as e:
        st.error(f"Error al generar el resumen: {str(e)}")
        return "Error al generar el resumen."

# Función para validar el nombre del archivo
def validate_file_type(filename, expected_type):
    return expected_type.lower() in filename.lower()


# Título principal - agregar antes de los file uploaders
st.markdown("<h1 style='text-align: center; color: #FFFF;'>Resúmenes pliegos Solutia</h1>", unsafe_allow_html=True)

# Subir PPT
st.markdown("### **Sube tu PPT**")
ppt_file = st.file_uploader(
    "",  # Label vacío porque usamos markdown arriba
    type=["pdf"],
    key="ppt_uploader",
    help="Arrastra o selecciona tu archivo PPT aquí - Límite 200MB por archivo • PDF"
)

# Verificar si el archivo fue removido
if not ppt_file:
    if st.session_state.processed_files["ppt"] is not None:
        st.session_state.processed_files["ppt"] = None
        st.session_state.processed_summaries["ppt"] = None
elif ppt_file and validate_file_type(ppt_file.name, "ppt") or validate_file_type(ppt_file.name, "Pliego de Prescripciones Técnicas"):
    with st.spinner("Procesando PPT..."):
        ppt_text = extract_text_with_langchain(ppt_file)
        if ppt_text:
            st.session_state.processed_files["ppt"] = ppt_text
            if "ppt" not in st.session_state.display_order:
                ppt_summary = process_full_document(ppt_text, llm, task="Resumen de PPT")
                st.session_state.processed_summaries["ppt"] = ppt_summary
                st.session_state.display_order.insert(0, "ppt")
            st.success(f"PPT procesado correctamente: {ppt_file.name}")
else:
    st.error("El archivo subido no parece ser un PPT. Por favor, verifica el nombre del archivo.")
    st.session_state.processed_files["ppt"] = None

# Subir PCAP
st.markdown("### **Sube tu PCAP**")
pcap_file = st.file_uploader(
    "",  # Label vacío porque usamos markdown arriba
    type=["pdf"],
    key="pcap_uploader",
    help="Arrastra o selecciona tu archivo PCAP aquí - Límite 200MB por archivo • PDF"
)

# Verificar si el archivo fue removido
if not pcap_file:
    if st.session_state.processed_files["pcap"] is not None:
        st.session_state.processed_files["pcap"] = None
        st.session_state.processed_summaries["pcap"] = None
elif pcap_file and validate_file_type(pcap_file.name, "pcap") or validate_file_type(pcap_file.name, "Pliego de Cláusula Administrativa"):
    with st.spinner("Procesando PCAP..."):
        pcap_text = extract_text_with_langchain(pcap_file)
        if pcap_text:
            st.session_state.processed_files["pcap"] = pcap_text
            if "pcap" not in st.session_state.display_order:
                pcap_summary = process_full_document(pcap_text, llm, task="Resumen de PCAP")
                st.session_state.processed_summaries["pcap"] = pcap_summary
                st.session_state.display_order.insert(0, "pcap")
            st.success(f"PCAP procesado correctamente: {pcap_file.name}")
else:
    st.error("El archivo subido no parece ser un PCAP. Por favor, verifica el nombre del archivo.")
    st.session_state.processed_files["pcap"] = None

def show_summary(doc_type):
    if st.session_state.processed_summaries[doc_type]:
        st.markdown(f"### Resumen {doc_type.upper()}")
        st.markdown(st.session_state.processed_summaries[doc_type], unsafe_allow_html=True)
        doc_path = create_word_document_with_clean_formatting(st.session_state.processed_summaries[doc_type])
        with open(doc_path, "rb") as file:
            st.download_button(
                label=f"Descargar resumen {doc_type.upper()} en Word",
                data=file,
                file_name=f"resumen_{doc_type}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{doc_type}_1"
            )
        os.remove(doc_path)

# Mostrar resúmenes en el orden actual
summary_container = st.container()
with summary_container:
    for doc_type in st.session_state.display_order:
        show_summary(doc_type)

# Modificar el botón de generar nuevo resumen
if st.button("Generar nuevo resumen"):
    current_type = st.session_state.display_order[0] if st.session_state.display_order else None
    
    if current_type:
        with st.spinner(f"Regenerando resumen {current_type.upper()}..."):
            # Limpiar el archivo Word anterior si existe
            if os.path.exists("resumen_generado.docx"):
                try:
                    os.remove("resumen_generado.docx")
                except:
                    pass
            
            # Generar nuevo resumen
            new_summary = process_full_document(
                st.session_state.processed_files[current_type],
                llm,
                task=f"Resumen de {current_type.upper()}"
            )
            
            # Actualizar el estado
            st.session_state.processed_summaries[current_type] = new_summary
            
            # Limpiar y actualizar el orden de visualización
            st.session_state.display_order = [current_type]
            
            st.success(f"Resumen {current_type.upper()} regenerado")
            st.rerun()

